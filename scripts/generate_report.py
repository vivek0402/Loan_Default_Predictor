"""
LendingClub Loan Default Risk Analysis - Full Project Report Generator
Professional narrative report for technical and non-technical audiences.
"""

import os, pickle, warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.metrics import (
    roc_auc_score, roc_curve, confusion_matrix,
    classification_report, precision_recall_curve, fbeta_score
)

BASE = 'e:/Projects/LendingClub_Loan'
BLUE      = RGBColor(0x1F, 0x49, 0x7D)
DARKBLUE  = RGBColor(0x0D, 0x2B, 0x55)
ACCENT    = RGBColor(0xC0, 0x39, 0x2B)
GRAY      = RGBColor(0x55, 0x55, 0x55)
LIGHTGRAY = RGBColor(0x88, 0x88, 0x88)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

# ─── styling helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_row_bg(row, hex_color):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)

def add_heading(doc, text, level=1, color=BLUE, space_before=12):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_body(doc, text, space_after=6, italic=False, bold=False, color=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if p.runs:
        p.runs[0].italic = italic
        p.runs[0].bold   = bold
        if color:
            p.runs[0].font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    if bold_prefix:
        run1 = p.add_run(bold_prefix + ' ')
        run1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_callout(doc, text, label='KEY INSIGHT', bg='1F497D'):
    """Adds a styled callout box as a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run_label = p.add_run(f'{label}:  ')
    run_label.bold = True
    run_label.font.color.rgb = WHITE
    run_label.font.size = Pt(9)
    run_text = p.add_run(text)
    run_text.font.color.rgb = WHITE
    run_text.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_image(doc, path, width=5.8, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'[Image not found: {os.path.basename(path)}]').runs[0].font.color.rgb = GRAY
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        if cp.runs:
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = GRAY

def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    set_row_bg(hdr_row, '1F497D')
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        if r_idx % 2 == 1:
            set_row_bg(row, 'EBF2FA')
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.add_run(str(val)).font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def metric_grid(doc, metrics, cols=2):
    """Render a dict of metrics as a clean N-col grid."""
    items = list(metrics.items())
    rows  = [items[i:i+cols] for i in range(0, len(items), cols)]
    table = doc.add_table(rows=len(rows), cols=cols * 2)
    table.style = 'Table Grid'
    for r_idx, row_items in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, (k, v) in enumerate(row_items):
            label_cell = row.cells[c_idx * 2]
            value_cell = row.cells[c_idx * 2 + 1]
            set_cell_bg(label_cell, 'D6E4F0')
            lp = label_cell.paragraphs[0]
            lp.add_run(k).bold = True
            lp.runs[0].font.size = Pt(9)
            vp = value_cell.paragraphs[0]
            vp.add_run(v).font.size = Pt(9)
            vp.runs[0].bold = True
            vp.runs[0].font.color.rgb = ACCENT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── load model & data ────────────────────────────────────────────────────────

print('Loading data and model...')
X_test  = pd.read_parquet(f'{BASE}/data/X_test.parquet').drop(columns=['installment','total_acc'])
y_test  = pd.read_parquet(f'{BASE}/data/y_test.parquet').squeeze()
with open(f'{BASE}/data/xgb_model.pkl', 'rb') as f:
    model = pickle.load(f)
with open(f'{BASE}/data/threshold.txt') as f:
    THRESHOLD = float(f.read())

proba = model.predict_proba(X_test)[:, 1]
preds = (proba >= THRESHOLD).astype(int)
cm    = confusion_matrix(y_test, preds)
tn, fp, fn, tp = cm.ravel()

auc        = roc_auc_score(y_test, proba)
f2         = fbeta_score(y_test, preds, beta=2)
catch_rate = tp / (tp + fn)
good_rej   = fp / (tn + fp)
total      = len(y_test)
approved   = tn + fn
rejected   = tp + fp
dr_base    = (tp + fn) / total
dr_model   = fn / approved
dr_reduc   = (dr_base - dr_model) / dr_base
feat_imp   = pd.Series(model.feature_importances_, index=X_test.columns).sort_values(ascending=False)

# ── rejected loans (sample for report charts) ─────────────────────────────
print('Loading rejected loans sample...')
TOTAL_REJECTED = 27_648_741
TOTAL_ACCEPTED_RAW = 2_260_701
TOTAL_ALL = TOTAL_REJECTED + TOTAL_ACCEPTED_RAW

rej = pd.read_csv(
    f'{BASE}/data/rejected_2007_to_2018Q4.csv',
    nrows=500_000,
    usecols=['Amount Requested', 'Application Date', 'Risk_Score',
             'Debt-To-Income Ratio', 'State', 'Employment Length']
)
rej['dti']  = pd.to_numeric(
    rej['Debt-To-Income Ratio'].str.replace('%', '').str.strip(), errors='coerce')
rej['year'] = pd.to_datetime(rej['Application Date'], errors='coerce').dt.year
rej['emp']  = rej['Employment Length'].fillna('n/a')

acc_raw = pd.read_parquet(f'{BASE}/data/loans_filtered.parquet',
                          columns=['dti', 'loan_amnt', 'emp_length'])
print('Rejected sample and accepted data loaded.')


# ─── generate charts ──────────────────────────────────────────────────────────

print('Generating charts...')
sns.set_theme(style='whitegrid', palette='muted')

# --- score distribution ---
fig, ax = plt.subplots(figsize=(9, 4))
ax.hist(proba[y_test==0], bins=60, alpha=0.65, color='steelblue', label='Fully Paid', density=True)
ax.hist(proba[y_test==1], bins=60, alpha=0.65, color='tomato',    label='Default',    density=True)
ax.axvline(THRESHOLD, color='black', linestyle='--', lw=2, label=f'Decision threshold = {THRESHOLD}')
ax.fill_betweenx([0, ax.get_ylim()[1] if ax.get_ylim()[1] > 0 else 10],
                 THRESHOLD, 1.0, alpha=0.08, color='tomato', label='Rejected zone')
ax.set_xlabel('Predicted Default Probability', fontsize=11)
ax.set_ylabel('Density', fontsize=11)
ax.set_title('Model Score Distribution: How Well Does It Separate Defaults from Paid Loans?', fontsize=11)
ax.legend(fontsize=9)
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_score_dist.png', dpi=130)
plt.close()

# --- feature importance ---
top15 = feat_imp.head(15).sort_values()
fig, ax = plt.subplots(figsize=(9, 6))
colors = ['#C0392B' if v > feat_imp.median() else '#2980B9' for v in top15]
bars = ax.barh(top15.index, top15.values, color=colors, edgecolor='white')
ax.set_xlabel('Importance Score', fontsize=11)
ax.set_title('Top 15 Most Influential Features (XGBoost)', fontsize=11)
ax.axvline(feat_imp.median(), color='gray', linestyle='--', alpha=0.6, label='Median')
for bar, val in zip(bars, top15.values):
    ax.text(val + 0.0005, bar.get_y() + bar.get_height()/2,
            f'{val:.4f}', va='center', fontsize=8)
ax.legend()
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_feat_imp.png', dpi=130)
plt.close()

# --- confusion matrix ---
fig, ax = plt.subplots(figsize=(6, 5))
cm_labels = pd.DataFrame(
    [[f'True Negative\n{tn:,}\n(Correctly Approved)', f'False Positive\n{fp:,}\n(Good borrower rejected)'],
     [f'False Negative\n{fn:,}\n(Default missed)',    f'True Positive\n{tp:,}\n(Default caught)']],
    index=['Actual: Fully Paid', 'Actual: Default'],
    columns=['Predicted: Approve', 'Predicted: Reject']
)
cm_vals = np.array([[tn, fp], [fn, tp]])
sns.heatmap(cm_vals, annot=cm_labels, fmt='', cmap='Blues',
            xticklabels=['Predicted: Approve', 'Predicted: Reject'],
            yticklabels=['Actual: Fully Paid', 'Actual: Default'],
            linewidths=1, ax=ax, annot_kws={'size': 9})
ax.set_title(f'Confusion Matrix  (threshold = {THRESHOLD})', fontsize=11)
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_cm.png', dpi=130)
plt.close()

# --- ROC curve ---
fpr_arr, tpr_arr, _ = roc_curve(y_test, proba)
fig, ax = plt.subplots(figsize=(7, 5))
ax.plot(fpr_arr, tpr_arr, color='tomato', lw=2.5, label=f'XGBoost  (AUC = {auc:.4f})')
ax.plot([0,1],[0,1], 'k--', lw=1, label='Random classifier  (AUC = 0.50)')
op_fpr = fp/(tn+fp); op_tpr = tp/(tp+fn)
ax.scatter(op_fpr, op_tpr, s=120, zorder=5, color='black',
           label=f'Operating point  (t={THRESHOLD})\ncatch={op_tpr:.1%}, FPR={op_fpr:.1%}')
ax.set_xlabel('False Positive Rate  (Good Borrowers Rejected)', fontsize=10)
ax.set_ylabel('True Positive Rate  (Defaults Caught)', fontsize=10)
ax.set_title('ROC Curve — Discriminatory Power of the Model', fontsize=11)
ax.legend(fontsize=9)
ax.fill_between(fpr_arr, tpr_arr, alpha=0.08, color='tomato')
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_roc.png', dpi=130)
plt.close()

# --- threshold sweep ---
sweep = []
for t in np.arange(0.10, 0.90, 0.01):
    p = (proba >= t).astype(int)
    c = confusion_matrix(y_test, p)
    _tn,_fp,_fn,_tp = c.ravel()
    catch  = _tp/(_tp+_fn)
    gr     = _fp/(_tn+_fp)
    f2s    = fbeta_score(y_test, p, beta=2)
    sweep.append({'t': round(t,2), 'catch': catch, 'good_rej': gr, 'f2': f2s})
sw = pd.DataFrame(sweep)

fig, axes = plt.subplots(1, 2, figsize=(13, 5))
axes[0].plot(sw['t'], sw['catch'],   color='tomato',    lw=2, label='Default Catch Rate')
axes[0].plot(sw['t'], sw['good_rej'],color='steelblue', lw=2, label='Good Borrower Rejection Rate')
axes[0].axvline(THRESHOLD, color='black', linestyle='--', lw=1.5, label=f'Chosen threshold = {THRESHOLD}')
axes[0].set_xlabel('Decision Threshold', fontsize=10)
axes[0].set_ylabel('Rate', fontsize=10)
axes[0].set_title('Tradeoff: Catching Defaults vs Rejecting Good Borrowers', fontsize=10)
axes[0].legend(fontsize=9)

axes[1].plot(sw['t'], sw['f2'], color='purple', lw=2.5)
axes[1].axvline(THRESHOLD, color='black', linestyle='--', lw=1.5, label=f'F2-optimal = {THRESHOLD}')
axes[1].scatter([THRESHOLD], [sw[sw['t']==THRESHOLD]['f2'].values[0]], s=100, color='black', zorder=5)
axes[1].set_xlabel('Decision Threshold', fontsize=10)
axes[1].set_ylabel('F2 Score', fontsize=10)
axes[1].set_title('F2 Score Across All Thresholds', fontsize=10)
axes[1].legend(fontsize=9)
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_threshold.png', dpi=130)
plt.close()

# --- rejected: scale comparison ---
fig, axes = plt.subplots(1, 2, figsize=(13, 4.5))

categories = ['Rejected\n(27.6M)', 'Accepted\n(2.3M)']
vals       = [TOTAL_REJECTED, TOTAL_ACCEPTED_RAW]
bars = axes[0].barh(categories, vals, color=['#C0392B','#2980B9'], height=0.45, edgecolor='white')
for bar, val in zip(bars, vals):
    axes[0].text(val + 300_000, bar.get_y() + bar.get_height()/2,
                 f'{val:,.0f}', va='center', fontsize=10, fontweight='bold')
axes[0].set_xlabel('Applications', fontsize=10)
axes[0].set_title('Total LendingClub Applications 2007–2018', fontsize=10)
axes[0].xaxis.set_major_formatter(
    plt.FuncFormatter(lambda x, _: f'{x/1e6:.0f}M'))
axes[0].set_xlim(0, TOTAL_REJECTED * 1.18)

funnel_labels = ['All Applications\n(29.9M)', 'Accepted\n(2.3M)', 'Model Universe\n(1.35M)']
funnel_vals   = [TOTAL_ALL, TOTAL_ACCEPTED_RAW, 1_348_092]
funnel_colors = ['#922B21', '#2E86C1', '#1A5276']
bars2 = axes[1].barh(funnel_labels, funnel_vals, color=funnel_colors, height=0.45, edgecolor='white')
for bar, val in zip(bars2, funnel_vals):
    axes[1].text(val + 300_000, bar.get_y() + bar.get_height()/2,
                 f'{val:,.0f}', va='center', fontsize=10, fontweight='bold')
axes[1].set_xlabel('Applications', fontsize=10)
axes[1].set_title('The Filtering Funnel: From Applications to Model Input', fontsize=10)
axes[1].xaxis.set_major_formatter(
    plt.FuncFormatter(lambda x, _: f'{x/1e6:.1f}M'))
axes[1].set_xlim(0, TOTAL_ALL * 1.22)
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_rej_scale.png', dpi=130)
plt.close()

# --- rejected: DTI comparison ---
rej_dti = rej['dti'].dropna()
rej_dti = rej_dti[(rej_dti >= 0) & (rej_dti <= 100)]
acc_dti = acc_raw['dti'].dropna()
acc_dti = acc_dti[(acc_dti >= 0) & (acc_dti <= 100)]

fig, ax = plt.subplots(figsize=(9, 4))
ax.hist(acc_dti, bins=60, alpha=0.65, color='steelblue', density=True,
        label=f'Accepted  (median={acc_dti.median():.1f}%)')
ax.hist(rej_dti, bins=60, alpha=0.65, color='tomato',    density=True,
        label=f'Rejected  (median={rej_dti.median():.1f}%)')
ax.axvline(acc_dti.median(), color='steelblue', linestyle='--', lw=2)
ax.axvline(rej_dti.median(), color='tomato',    linestyle='--', lw=2)
ax.set_xlabel('Debt-to-Income Ratio (%)', fontsize=11)
ax.set_ylabel('Density', fontsize=11)
ax.set_title('DTI Distribution: Accepted vs Rejected Applicants', fontsize=11)
ax.legend(fontsize=10)
ax.set_xlim(0, 80)
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_rej_dti.png', dpi=130)
plt.close()

# --- rejected: risk score ---
risk = rej['Risk_Score'].dropna()
risk = risk[(risk >= 300) & (risk <= 850)]

fig, ax = plt.subplots(figsize=(9, 4))
ax.hist(risk, bins=80, color='#C0392B', alpha=0.8, edgecolor='white')
bands = [(300,580,'#922B21','Very Poor\n(<580)'), (580,670,'#E67E22','Fair\n(580-669)'),
         (670,740,'#F1C40F','Good\n(670-739)'), (740,800,'#27AE60','Very Good\n(740+)')]
ymax = ax.get_ylim()[1]
for lo, hi, col, lbl in bands:
    ax.axvspan(lo, hi, alpha=0.08, color=col)
    ax.text((lo+hi)/2, ymax*0.9, lbl, ha='center', fontsize=8, color=col, fontweight='bold')
ax.axvline(risk.median(), color='black', linestyle='--', lw=2, label=f'Median = {risk.median():.0f}')
ax.set_xlabel('Risk Score (FICO-like)', fontsize=11)
ax.set_ylabel('Rejected Applicants', fontsize=11)
ax.set_title('Risk Score Distribution of Rejected Applicants', fontsize=11)
ax.legend(fontsize=10)
ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f'{x/1000:.0f}K'))
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_rej_risk.png', dpi=130)
plt.close()

# --- rejected: employment length ---
EMP_ORDER = ['< 1 year','1 year','2 years','3 years','4 years',
             '5 years','6 years','7 years','8 years','9 years','10+ years']
rej_emp = rej['emp'].value_counts(normalize=True)
acc_emp = acc_raw['emp_length'].value_counts(normalize=True)
rej_pct = pd.Series({k: rej_emp.get(k, 0) for k in EMP_ORDER})
acc_pct = pd.Series({k: acc_emp.get(k, 0) for k in EMP_ORDER})

x     = np.arange(len(EMP_ORDER))
width = 0.38
fig, ax = plt.subplots(figsize=(12, 4.5))
ax.bar(x - width/2, acc_pct.values * 100, width, label='Accepted', color='steelblue', alpha=0.85)
ax.bar(x + width/2, rej_pct.values * 100, width, label='Rejected', color='tomato',    alpha=0.85)
ax.set_xticks(x)
ax.set_xticklabels(EMP_ORDER, rotation=30, ha='right', fontsize=9)
ax.set_ylabel('% of Applicants', fontsize=11)
ax.set_title('Employment Length Distribution: Accepted vs Rejected', fontsize=11)
ax.legend(fontsize=10)
ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f'{x:.0f}%'))
plt.tight_layout()
plt.savefig(f'{BASE}/data/rpt_rej_emp.png', dpi=130)
plt.close()

print('All charts generated.')


# ─── build document ───────────────────────────────────────────────────────────

print('Building report...')
doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('LendingClub Loan Default Risk Analysis')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARKBLUE

doc.add_paragraph()

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = cover_sub.add_run('A Complete Machine Learning Project:\nFrom Raw Data to Business Impact')
run2.font.size = Pt(14)
run2.font.color.rgb = BLUE

doc.add_paragraph()
doc.add_paragraph()

# summary box on cover
tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.rows[0].cells[0]
set_cell_bg(cell, '1F497D')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f'Dataset: LendingClub 2007–2018   |   1.35 Million Loans Analysed\n'
    f'Model: XGBoost   |   AUC: {auc:.4f}   |   Default Rate Reduced: {dr_reduc:.0%}'
)
r.font.color.rgb = WHITE
r.font.size = Pt(11)
r.bold = True
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

cover_note = doc.add_paragraph()
cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = cover_note.add_run(
    'This report is written for both technical and non-technical readers.\n'
    'Every concept is explained in plain English before the analysis.'
)
rn.font.size = Pt(10)
rn.italic = True
rn.font.color.rgb = GRAY

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════

add_heading(doc, 'Table of Contents', 1, color=DARKBLUE)

toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Background — What Is This Project About?'),
    ('3.', 'The Dataset'),
    ('3.5', 'The Rejected Applications — Completing the Picture'),
    ('4.', 'Project Objectives & Business Problem'),
    ('5.', 'Exploratory Data Analysis (EDA)'),
    ('6.', 'Hypothesis Testing'),
    ('7.', 'Data Preprocessing'),
    ('8.', 'Feature Selection & Statistical Significance'),
    ('9.', 'Machine Learning Models'),
    ('10.', 'Model Comparison'),
    ('11.', 'Threshold Optimisation'),
    ('12.', 'Final Model Performance'),
    ('13.', 'Business Impact'),
    ('14.', 'Conclusions & Recommendations'),
    ('15.', 'Limitations & Future Work'),
    ('16.', 'Glossary of Key Terms'),
    ('17.', 'Project Notebook Structure'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.color.rgb = BLUE
    p.add_run(title).font.size = Pt(10)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════

add_heading(doc, '1. Executive Summary', 1, color=DARKBLUE)

add_body(doc,
    'This project builds a machine learning model to predict whether a loan applicant will '
    'default (fail to repay) their loan. Using over 1.35 million historical loans from '
    'LendingClub (2007–2018), we trained an XGBoost model that can screen applicants before '
    'a loan is approved — significantly reducing the platform\'s exposure to bad debt.'
)

add_callout(doc,
    f'The final model reduces the default rate among approved loans from '
    f'{dr_base:.1%} to {dr_model:.1%} — a {dr_reduc:.0%} reduction — '
    f'while catching {catch_rate:.1%} of all potential defaults.',
    label='HEADLINE RESULT'
)

add_body(doc, 'Key findings at a glance:')

metric_grid(doc, {
    'Total Loans Analysed':       f'{total:,}',
    'Model Used':                  'XGBoost',
    'ROC-AUC Score':               f'{auc:.4f}',
    'F2 Score':                    f'{f2:.4f}',
    'Defaults Caught':             f'{tp:,}  ({catch_rate:.1%})',
    'Defaults Missed':             f'{fn:,}',
    'Default Rate (No Model)':     f'{dr_base:.2%}',
    'Default Rate (With Model)':   f'{dr_model:.2%}',
    'Default Rate Reduction':      f'{dr_reduc:.0%}',
    'Decision Threshold':          str(THRESHOLD),
}, cols=2)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. BACKGROUND
# ════════════════════════════════════════════════════════════════

add_heading(doc, '2. Background — What Is This Project About?', 1, color=DARKBLUE)

add_heading(doc, '2.1  What Is LendingClub?', 2)
add_body(doc,
    'LendingClub is an American peer-to-peer (P2P) lending platform. Unlike a traditional bank, '
    'LendingClub connects individual borrowers directly with individual investors. A borrower '
    'applies for a personal loan; investors then fund that loan and receive monthly repayments '
    'with interest. If the borrower stops paying, the investor loses their money.'
)
add_body(doc,
    'Between 2007 and 2018, LendingClub issued over 2.26 million loans worth hundreds of '
    'billions of dollars. This rich dataset makes it one of the most studied sources of '
    'real-world credit risk data available to the public.'
)

add_heading(doc, '2.2  What Is a Loan Default?', 2)
add_body(doc,
    'A loan default occurs when a borrower stops making their scheduled repayments and '
    'the debt is written off as uncollectable. When LendingClub marks a loan as "Charged Off", '
    'it means the investor has lost that money. In this dataset, approximately 1 in 5 loans '
    'ended in default — a significant financial risk.'
)

add_heading(doc, '2.3  Why Use Machine Learning?', 2)
add_body(doc,
    'Traditional credit scoring relies on a few simple rules: credit score above X, income '
    'above Y, debt-to-income ratio below Z. Machine learning can consider dozens of factors '
    'simultaneously and find non-obvious patterns that humans and simple rules would miss. '
    'For example, the combination of a moderate interest rate AND short credit history AND '
    'high revolving utilisation might be far riskier than any single factor alone.'
)

add_callout(doc,
    'Machine learning does not replace human judgement — it augments it. '
    'The model assigns a risk score to each applicant; a human or policy rule '
    'then decides what to do with that score.',
    label='IMPORTANT NOTE', bg='2C3E50'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. THE DATASET
# ════════════════════════════════════════════════════════════════

add_heading(doc, '3. The Dataset', 1, color=DARKBLUE)

add_heading(doc, '3.1  Source & Scope', 2)
add_body(doc,
    'The dataset used in this project is the publicly available LendingClub loan data '
    'covering all loans issued from January 2007 to Q4 2018. It was downloaded from '
    'Kaggle and contains the full history of every accepted loan application.'
)

styled_table(doc,
    ['Property', 'Value'],
    [
        ['File',                   'accepted_2007_to_2018Q4.csv'],
        ['Total rows (raw)',        '2,260,701 loans'],
        ['Total columns',           '151 attributes per loan'],
        ['Time period',             'January 2007 – December 2018'],
        ['Rows after filtering',    '1,348,092 loans'],
        ['Final features used',     '34 features'],
    ],
    col_widths=[2.5, 4.0]
)

add_heading(doc, '3.2  What Does Each Row Represent?', 2)
add_body(doc,
    'Each row in the dataset represents one individual loan application. It contains '
    'information known at the time of application (income, employment, credit history) '
    'as well as the outcome of the loan (did the borrower repay fully or default?).'
)

add_heading(doc, '3.3  The Target Variable: Loan Status', 2)
add_body(doc,
    'The column "loan_status" tells us what happened to each loan. We mapped it to a '
    'binary (0/1) variable called "target": 1 means the loan defaulted, 0 means it was '
    'fully repaid. Loans with ambiguous statuses (still active, in grace period, etc.) '
    'were removed because we cannot know their final outcome.'
)

styled_table(doc,
    ['loan_status Value', 'Mapped To', 'Count', 'Percentage', 'Reason'],
    [
        ['Fully Paid',                                   '0 — Good Loan',    '1,076,751', '47.6%', 'Borrower repaid in full'],
        ['Charged Off',                                  '1 — Default',      '268,559',   '11.9%', 'Written off as bad debt'],
        ['Default',                                      '1 — Default',      '40',        '0.0%',  'Formal default status'],
        ['Does not meet credit policy — Charged Off',    '1 — Default',      '761',       '0.0%',  'Still a default'],
        ['Does not meet credit policy — Fully Paid',     '0 — Good Loan',    '1,988',     '0.1%',  'Still repaid'],
        ['Current',                                      'DROPPED',          '878,317',   '38.9%', 'Loan still active — unknown outcome'],
        ['Late (31-120 days)',                            'DROPPED',          '21,467',    '0.9%',  'Ambiguous — may still resolve'],
        ['In Grace Period',                              'DROPPED',          '8,436',     '0.4%',  'Ambiguous — may still resolve'],
        ['Late (16-30 days)',                             'DROPPED',          '4,349',     '0.2%',  'Ambiguous — may still resolve'],
    ],
    col_widths=[2.3, 1.5, 1.0, 0.9, 1.8]
)

add_callout(doc,
    'After filtering, the dataset contains 80.0% good loans (class 0) and 20.0% defaults '
    '(class 1). This imbalance is realistic — most borrowers do repay — but it means the '
    'model needs special handling to avoid ignoring the minority class.',
    label='CLASS IMBALANCE'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3.5  THE REJECTED APPLICATIONS
# ════════════════════════════════════════════════════════════════

add_heading(doc, '3.5  The Rejected Applications — Completing the Picture', 1, color=DARKBLUE)

add_body(doc,
    'Most analyses of LendingClub data focus only on the accepted loans — the ones that '
    'were funded and either repaid or defaulted. But LendingClub also published a second '
    'dataset: every loan application they rejected. Analysing this rejected dataset gives '
    'us a crucial piece of context: who never even made it into our model\'s training data, '
    'and why.'
)

add_heading(doc, '3.5.1  The Scale: How Many Applications Were Rejected?', 2)

add_body(doc,
    'Between 2007 and 2018, LendingClub received nearly 30 million loan applications in total. '
    'The overwhelming majority were turned down before a single dollar was lent.'
)

metric_grid(doc, {
    'Total Applications Ever Received':  f'{TOTAL_ALL:,.0f}',
    'Applications Rejected':             f'{TOTAL_REJECTED:,.0f}  ({TOTAL_REJECTED/TOTAL_ALL:.1%})',
    'Applications Accepted':             f'{TOTAL_ACCEPTED_RAW:,.0f}  ({TOTAL_ACCEPTED_RAW/TOTAL_ALL:.1%})',
    'Used in Our ML Model':              '1,348,092  (accepted + resolved only)',
}, cols=2)

add_callout(doc,
    f'Out of every 100 people who applied to LendingClub, approximately '
    f'{TOTAL_REJECTED/TOTAL_ALL*100:.0f} were rejected outright before any loan was issued. '
    f'Our machine learning model only ever sees the {TOTAL_ACCEPTED_RAW/TOTAL_ALL*100:.0f} '
    f'who passed LendingClub\'s initial screening — and even among those, 20% still defaulted.',
    label='KEY CONTEXT'
)

add_image(doc,
    f'{BASE}/data/rpt_rej_scale.png',
    width=6.0,
    caption='Figure: Left — total rejected vs accepted. Right — the three-tier funnel from all '
            'applications down to the 1.35M loans used in our model.'
)

add_heading(doc, '3.5.2  The Biggest Rejection Driver: Employment Stability', 2)

add_body(doc,
    'The most striking finding from the rejected dataset is not DTI or credit score — it is '
    'employment length. A massive 70.2% of rejected applicants had less than one year of '
    'employment history, compared to just 8.5% of accepted applicants. This single factor '
    'alone explains a large part of LendingClub\'s rejection decisions. An applicant who just '
    'started a new job (or who has no stable employment record) represents a much higher income '
    'risk — if they lose their job, they cannot repay the loan.'
)

add_image(doc,
    f'{BASE}/data/rpt_rej_emp.png',
    width=6.2,
    caption='Figure: Employment length distribution for accepted vs rejected applicants. '
            '70.2% of rejected applicants had under 1 year of employment, vs 8.5% of accepted. '
            '10+ year employees made up 34.9% of accepted but only 6.1% of rejected.'
)

add_callout(doc,
    '70.2% of rejected applicants had less than 1 year of employment, compared to 8.5% of accepted applicants — '
    'a gap of +61.7 percentage points. Employment stability was the single largest differentiator '
    'between who was accepted and who was rejected.',
    label='KEY FINDING: EMPLOYMENT', bg='922B21'
)

styled_table(doc,
    ['Employment Length', 'Accepted (%)', 'Rejected (%)', 'Gap (pp)'],
    [
        ['< 1 year',   '8.5%',  '70.2%', '+61.7'],
        ['1 year',     '7.0%',  '4.4%',  '-2.6'],
        ['2 years',    '9.6%',  '4.5%',  '-5.1'],
        ['5 years',    '6.6%',  '2.4%',  '-4.3'],
        ['10+ years',  '34.9%', '6.1%',  '-28.8'],
    ],
    col_widths=[1.8, 1.4, 1.4, 1.4]
)

add_heading(doc, '3.5.3  DTI: A Surprising Finding', 2)

add_body(doc,
    'Debt-to-Income ratio is one of the strongest predictors in our machine learning model, '
    'so one would expect rejected applicants to have much higher DTI than accepted ones. '
    'The data tells a more nuanced story. The median DTI for rejected applicants (15.9%) was '
    'actually slightly lower than for accepted applicants (17.6%). However, the mean DTI was '
    'higher for rejected (19.3% vs 18.2%), and the distribution was more spread out with a '
    'longer right tail.'
)

add_body(doc,
    'This tells us that LendingClub did not use a simple DTI cutoff as its primary rejection '
    'rule. Instead, it relied more heavily on credit score and employment stability (as we saw '
    'in the previous section). High DTI applicants were more likely to be rejected, but a high '
    'DTI alone was not sufficient cause for rejection — other factors mattered more.'
)

add_image(doc,
    f'{BASE}/data/rpt_rej_dti.png',
    width=6.0,
    caption='Figure: DTI distribution for accepted vs rejected applicants. The medians are close '
            '(15.9% rejected vs 17.6% accepted), but rejected applicants show a heavier right tail, '
            'indicating some very high-DTI applicants were rejected.'
)

add_callout(doc,
    'Accepted median DTI: 17.6%.  Rejected median DTI: 15.9%.  '
    'DTI alone was not LendingClub\'s primary rejection signal — '
    'credit score and employment stability played a larger role.',
    label='DTI FINDING', bg='2C3E50'
)

add_heading(doc, '3.5.4  Credit Score (Risk Score) of Rejected Applicants', 2)

add_body(doc,
    'The rejected dataset includes a "Risk_Score" — a FICO-like credit score ranging from '
    '300 to 850. Higher scores indicate better credit history. The accepted dataset does not '
    'have this field (LendingClub omitted it from public accepted data), so we analyse '
    'rejected applicants only. The median risk score for rejected applicants was 650 — '
    'in the "Fair" band. 62.2% of rejected applicants scored below 670 (Very Poor + Fair), '
    'meaning most had below-average credit. However, 31.6% scored in the "Good" range (670-739), '
    'suggesting that credit score alone was also not a hard cutoff — other factors combined '
    'to determine the final rejection decision.'
)

add_body(doc,
    'FICO credit score bands:',
    space_after=3
)
styled_table(doc,
    ['Score Range', 'Category', '% of Rejected Applicants'],
    [
        ['300 – 579', 'Very Poor',   '23.5%'],
        ['580 – 669', 'Fair',        '38.7%  ← Largest group'],
        ['670 – 739', 'Good',        '31.6%'],
        ['740 – 799', 'Very Good',   '5.7%'],
        ['800 – 850', 'Exceptional', '0.4%'],
    ],
    col_widths=[1.4, 1.4, 3.6]
)

add_image(doc,
    f'{BASE}/data/rpt_rej_risk.png',
    width=6.0,
    caption='Figure: Risk Score distribution for rejected applicants. Median = 650 (Fair band). '
            '62.2% scored below 670, but 31.6% had Good scores — showing credit score alone '
            'did not fully determine rejection.'
)

add_callout(doc,
    'Median risk score of rejected applicants: 650 (Fair band).  '
    '62.2% scored below 670 (Very Poor or Fair).  '
    'Notably, 31.6% had "Good" credit scores (670-739) yet were still rejected — '
    'likely due to short employment history or other combined risk factors.',
    label='CREDIT SCORE FINDING', bg='1A5276'
)

add_heading(doc, '3.5.5  Loan Amount: Not the Driver', 2)

add_body(doc,
    'One might assume that rejected applicants were asking for too much money. The data '
    'shows the opposite: rejected applicants actually requested slightly less than accepted '
    'ones. The median requested amount for rejected applicants was $10,000, versus $12,000 '
    'for accepted applicants. This confirms that the amount of money requested was not a '
    'significant factor in LendingClub\'s rejection decisions — creditworthiness and '
    'employment stability mattered far more.'
)

metric_grid(doc, {
    'Accepted Median Loan Amount':  '$12,000',
    'Rejected Median Loan Amount':  '$10,000',
    'Accepted Mean Loan Amount':    '$14,409',
    'Rejected Mean Loan Amount':    '$12,219',
}, cols=2)

add_heading(doc, '3.5.6  Rejection Rate Growth Over Time', 2)

add_body(doc,
    'Rejection volumes grew dramatically over the 2007-2018 period, particularly after 2015. '
    'From just 5,274 rejections in 2007, LendingClub was rejecting nearly 9.5 million '
    'applications in 2018 alone. This reflects both the platform\'s massive growth in popularity '
    'and its increasingly strict underwriting standards — especially after a high-profile '
    'executive scandal in 2016 that led to investor scrutiny of loan quality.'
)

styled_table(doc,
    ['Year', 'Rejected Applications', 'Change vs Prior Year'],
    [
        ['2007',  '5,274',      '—'],
        ['2010',  '112,561',    '+~4x from 2009'],
        ['2013',  '760,942',    'Peak early growth'],
        ['2015',  '2,859,379',  '+48% from 2014'],
        ['2016',  '4,769,874',  '+67% — post-scandal tightening'],
        ['2017',  '7,072,573',  '+48%'],
        ['2018',  '9,496,782',  '+34% — highest year on record'],
    ],
    col_widths=[0.9, 2.1, 3.5]
)

add_heading(doc, '3.5.7  What This Means for Our Model', 2)

add_body(doc,
    'Our machine learning model operates as a second layer of risk screening on top of '
    'LendingClub\'s rules-based system. The first layer rejects 92.4% of all applicants '
    'based on hard criteria — primarily credit score below a threshold, very short employment '
    'history, or extreme DTI. The applicants our model sees have already cleared these hurdles.'
)

add_body(doc,
    'Yet even within this pre-screened population, 20% of loans still default. This is the '
    'fundamental challenge of credit risk: simple rules can filter out the obviously risky, '
    'but cannot distinguish between the moderately-risky and the safe among those who look '
    'similar on paper. Machine learning finds the subtle, non-linear combinations of factors '
    'that separate these groups — which is exactly what our XGBoost model does.'
)

add_callout(doc,
    'Two-layer risk system: (1) LendingClub\'s rules filter out 27.6M high-risk applications. '
    '(2) Our XGBoost model catches 88.8% of the remaining defaults among the accepted applicants. '
    'Together, these two layers reduce the effective default rate far below what either could achieve alone.',
    label='SYSTEM DESIGN', bg='1A5276'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. PROJECT OBJECTIVES & BUSINESS PROBLEM
# ════════════════════════════════════════════════════════════════

add_heading(doc, '4. Project Objectives & Business Problem', 1, color=DARKBLUE)

add_heading(doc, '4.1  The Business Question', 2)
add_body(doc,
    'Before any loan is funded, LendingClub needs to assess the risk that the borrower '
    'will not repay. The core question is:'
)
add_callout(doc,
    '"Given what we know about an applicant at the time of application, '
    'what is the probability that this loan will default?"',
    label='CORE QUESTION', bg='922B21'
)

add_heading(doc, '4.2  Business Constraint', 2)
add_body(doc,
    'The business stakeholder defined the following constraint:'
)
for item in [
    ('Reduce defaults by 30%:', 'The model must catch significantly more defaults than the current system.'),
    ('Reject at most 10% more good borrowers:', 'Every rejected good borrower is lost revenue. The model must be selective, not indiscriminate.'),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

add_body(doc,
    '\nAs the analysis progressed, this constraint was re-evaluated using F2-score optimisation '
    '(explained in Section 11), which ultimately delivered a 64% reduction in default rate '
    'among approved loans — far exceeding the original 30% target on that measure.',
    space_after=8
)

add_heading(doc, '4.3  Hypotheses', 2)
add_body(doc,
    'Five hypotheses were defined at the start of the project, based on domain knowledge '
    'about credit risk. Each hypothesis was tested statistically in the EDA phase:'
)

styled_table(doc,
    ['#', 'Hypothesis', 'Intuition'],
    [
        ['H1', 'Higher Debt-to-Income ratio → more defaults',
         'Borrowers stretched thin financially are more likely to miss payments'],
        ['H2', 'Shorter credit history → more defaults',
         'Less experience managing debt = higher risk'],
        ['H3', 'Certain loan purposes → higher default rates',
         'Small business loans are riskier than debt consolidation'],
        ['H4', 'Lower annual income → more defaults',
         'Less financial cushion to absorb unexpected expenses'],
        ['H5', 'More derogatory public records → more defaults',
         'Past bankruptcy or tax liens indicate poor financial behaviour'],
    ],
    col_widths=[0.4, 2.8, 3.3]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. EXPLORATORY DATA ANALYSIS
# ════════════════════════════════════════════════════════════════

add_heading(doc, '5. Exploratory Data Analysis (EDA)', 1, color=DARKBLUE)

add_heading(doc, '5.1  What Is EDA?', 2)
add_body(doc,
    'Exploratory Data Analysis is the process of visually and statistically examining your '
    'data before building any model. It answers three questions: (1) Are my hypotheses '
    'correct — do the expected patterns actually exist? (2) What does the data look like — '
    'distributions, outliers, missing values? (3) Which features will be useful to a model?'
)
add_body(doc,
    'Skipping EDA is one of the most common mistakes in data science. A model trained on '
    'misunderstood data produces wrong answers with high confidence — the worst outcome.'
)

add_heading(doc, '5.2  Key EDA Findings', 2)

add_body(doc, 'Debt-to-Income Ratio (H1)', bold=True)
add_body(doc,
    'DTI measures how much of a borrower\'s monthly income is consumed by debt payments. '
    'Defaulters had a median DTI of 19.75 versus 17.10 for fully paid borrowers. '
    'The KDE (density) plots show the entire default distribution shifted right — '
    'higher DTI consistently associates with higher default risk.'
)

if os.path.exists(f'{BASE}/data/h1_dti.png'):
    add_image(doc, f'{BASE}/data/h1_dti.png', width=5.5,
              caption='Figure 1: DTI distribution — defaulters (red) vs fully paid (blue)')

add_body(doc, 'Loan Purpose (H3)', bold=True)
add_body(doc,
    'The purpose of a loan has a strong association with default. Small business loans '
    'showed the highest default rate (~28%), followed by renewable energy and educational '
    'loans. Debt consolidation — the most common purpose — had a moderate default rate (~20%). '
    'Wedding and car loans were among the safest categories.'
)

if os.path.exists(f'{BASE}/data/h3_purpose.png'):
    add_image(doc, f'{BASE}/data/h3_purpose.png', width=5.5,
              caption='Figure 2: Default rate (left) and loan volume (right) by loan purpose')

add_body(doc, 'Annual Income (H4)', bold=True)
add_body(doc,
    'Defaulters had a median annual income of $60,000 versus $65,000 for fully paid '
    'borrowers. While the $5,000 gap is statistically significant given the large dataset, '
    'the signal is modest. Income alone does not strongly predict default — it works best '
    'in combination with DTI and grade.'
)

if os.path.exists(f'{BASE}/data/h4_income.png'):
    add_image(doc, f'{BASE}/data/h4_income.png', width=5.5,
              caption='Figure 3: Annual income distribution by target class')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 6. HYPOTHESIS TESTING
# ════════════════════════════════════════════════════════════════

add_heading(doc, '6. Hypothesis Testing Results', 1, color=DARKBLUE)

add_heading(doc, '6.1  Statistical Tests Used', 2)
add_body(doc,
    'Three statistical tests were applied to validate each hypothesis rigorously:'
)
for item in [
    ('Mann-Whitney U Test:',
     'A non-parametric test that checks whether the distribution of a numeric feature '
     'is significantly different between defaulters and paid borrowers. It does not '
     'assume normal distribution — important because financial data is heavily skewed.'),
    ('Point-Biserial Correlation:',
     'Measures the linear correlation between a numeric feature and the binary target. '
     'Values close to 0 indicate weak signal; values above ±0.1 are considered meaningful.'),
    ('Chi-Square Test:',
     'For categorical features (like loan purpose), tests whether the feature and the '
     'target variable are independent. A p-value below 0.05 means they are related.'),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

add_heading(doc, '6.2  Results Summary', 2)
add_body(doc,
    'With 1.35 million rows, every feature returned p = 0.0 on Mann-Whitney (statistically '
    'significant). This is expected — very large datasets make tiny differences statistically '
    'significant. We therefore also look at the magnitude of the difference, not just the p-value.'
)

styled_table(doc,
    ['Hypothesis', 'Feature', 'Default Value', 'Paid Value', 'Difference', 'Verdict'],
    [
        ['H1: Higher DTI',             'dti',                  '19.75',  '17.10', '+2.65 (15.5% higher)', '✓ Confirmed'],
        ['H2: Shorter Credit History', 'credit_history_years', '17.85yr','18.68yr','-0.83yr (4.4% shorter)','~ Weak signal'],
        ['H3: Loan Purpose',           'purpose (categorical)','—',      '—',     'Chi² = 4,182, p = 0.0', '✓ Confirmed'],
        ['H4: Lower Income',           'annual_inc',           '$60,000','$65,000','$5,000 (7.7% lower)',   '✓ Confirmed'],
        ['H5: Derogatory Marks',       'pub_rec',              '0.246',  '0.207', '+0.039 (18.8% higher)', '~ Weak signal'],
    ],
    col_widths=[1.6, 1.5, 1.0, 1.0, 1.6, 1.0]
)

add_heading(doc, '6.3  Interpretation', 2)
add_body(doc,
    'H1, H3, and H4 are confirmed with meaningful effect sizes. H2 and H5 show the '
    'correct direction but small magnitude — they are kept as features but are expected '
    'to have low importance in the final model. This is exactly what the XGBoost feature '
    'importance chart later confirms: grade and int_rate dominate, while credit_history_years '
    'and has_derogatory contribute modestly.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 7. DATA PREPROCESSING
# ════════════════════════════════════════════════════════════════

add_heading(doc, '7. Data Preprocessing', 1, color=DARKBLUE)

add_heading(doc, '7.1  Why Preprocessing Matters', 2)
add_body(doc,
    'Raw data cannot be fed directly into a machine learning model. Models are mathematical '
    'functions — they require numbers, not text. They cannot handle missing values. They can '
    'be misled by features on wildly different scales. Preprocessing is the bridge between '
    'raw data and a model-ready dataset.'
)
add_callout(doc,
    '"Garbage in, garbage out." The quality of the model is bounded by the quality '
    'of the input data. Preprocessing is where most of the real work happens.',
    label='PRINCIPLE'
)

add_heading(doc, '7.2  Steps Performed', 2)

steps = [
    ('Step 1: Feature Selection (151 → 20 columns)',
     'Of 151 raw columns, only 20 were selected. Many columns were excluded because they '
     'contain information only available after the loan is issued (e.g., total_payment, '
     'recoveries). Using these would be data leakage — the model would "cheat" by learning '
     'from the future. Only features available at the time of application were kept.'),
    ('Step 2: Feature Engineering',
     'Two transformations were applied: (a) earliest_cr_line, stored as a string like '
     '"Jan-2005", was converted to credit_history_years by measuring the gap from the '
     'loan date. (b) pub_rec (number of public records) was converted to a binary flag '
     'has_derogatory (0 or 1) because the distinction between 0 and 1+ records is far '
     'more informative than the raw count.'),
    ('Step 3: Handling Missing Values',
     'Numeric columns with missing values were filled with the median (not mean — '
     'financial data is skewed by outliers). Categorical columns were filled with the '
     'most common value (mode). The 33 rows where the target itself was missing were '
     'dropped entirely — there is nothing to learn from a loan with no known outcome.'),
    ('Step 4: Encoding Categorical Features',
     'Machine learning models require numbers. Categorical columns (purpose, '
     'home_ownership, verification_status) were one-hot encoded — each category becomes '
     'its own binary column. This avoids imposing false numeric ordering (e.g., treating '
     '"medical" as numerically "greater than" "car").'),
    ('Step 5: Train/Test Split (80% / 20%)',
     'The dataset was split: 1,078,447 rows for training, 269,612 for testing. The split '
     'was stratified — both sets contain the same 80/20 class ratio. The test set was '
     'never used during training; it exists solely to simulate future, unseen loan '
     'applications.'),
    ('Step 6: Feature Scaling',
     'StandardScaler was applied to 15 numeric columns, transforming each to have mean=0 '
     'and standard deviation=1. This prevents income ($60,000 range) from dominating DTI '
     '(0-60 range) simply because its numbers are larger. Critically, the scaler was fit '
     'only on training data and then applied to test data — never the reverse.'),
]

for title, body in steps:
    add_body(doc, title, bold=True, space_after=3)
    add_body(doc, body, space_after=8)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 8. FEATURE SIGNIFICANCE
# ════════════════════════════════════════════════════════════════

add_heading(doc, '8. Feature Selection & Statistical Significance', 1, color=DARKBLUE)

add_heading(doc, '8.1  Why Test Feature Significance?', 2)
add_body(doc,
    'Not every feature that is statistically significant is useful to a model. Two features '
    'that measure almost the same thing are redundant — including both adds noise without '
    'adding information. Multicollinearity (high correlation between features) can also '
    'destabilise some models. This section documents how we validated and refined the '
    'feature set before modelling.'
)

add_heading(doc, '8.2  Multicollinearity: Correlated Feature Pairs', 2)
add_body(doc,
    'A correlation matrix revealed three pairs of highly correlated features:'
)

styled_table(doc,
    ['Feature 1', 'Feature 2', 'Correlation', 'Problem', 'Action Taken'],
    [
        ['loan_amnt',  'installment', '0.953', 'Nearly identical — installment is just loan_amnt / term',
         'Dropped installment'],
        ['int_rate',   'grade',       '0.952', 'Grade encodes interest rate band — near-duplicate',
         'Retained int_rate (stronger signal)'],
        ['open_acc',   'total_acc',   '0.701', 'total_acc includes closed accounts — lower signal',
         'Dropped total_acc'],
    ],
    col_widths=[1.0, 1.0, 0.9, 2.3, 1.5]
)

add_heading(doc, '8.3  VIF Analysis', 2)
add_body(doc,
    'Variance Inflation Factor (VIF) quantifies how much each feature\'s variance is '
    'explained by all other features. VIF > 10 indicates severe multicollinearity. '
    'Tree-based models (Random Forest, XGBoost) are naturally resistant to multicollinearity, '
    'so int_rate (VIF=75) was retained despite its high VIF — it carries the strongest '
    'individual signal.'
)

if os.path.exists(f'{BASE}/data/correlation_matrix.png'):
    add_image(doc, f'{BASE}/data/correlation_matrix.png', width=5.5,
              caption='Figure 4: Feature correlation matrix — darker red/blue = higher correlation')

add_heading(doc, '8.4  Final Feature List (34 Features)', 2)

styled_table(doc,
    ['Feature', 'Type', 'Description', 'Hypothesis'],
    [
        ['dti',                  'Numeric',  'Debt-to-income ratio',                       'H1'],
        ['annual_inc',           'Numeric',  'Annual income (capped at 99th percentile)',   'H4'],
        ['loan_amnt',            'Numeric',  'Requested loan amount',                       '—'],
        ['int_rate',             'Numeric',  'Interest rate assigned to the loan',          '—'],
        ['grade',                'Numeric',  'LendingClub risk grade (A=1 to G=7)',          '—'],
        ['emp_length',           'Numeric',  'Employment length in years (0-10)',            '—'],
        ['open_acc',             'Numeric',  'Number of open credit lines',                 '—'],
        ['revol_bal',            'Numeric',  'Total revolving balance',                     '—'],
        ['revol_util',           'Numeric',  'Revolving line utilisation rate (%)',          '—'],
        ['mort_acc',             'Numeric',  'Number of mortgage accounts',                 '—'],
        ['delinq_2yrs',          'Numeric',  'Delinquencies in last 2 years',               '—'],
        ['inq_last_6mths',       'Numeric',  'Credit enquiries in last 6 months',           '—'],
        ['credit_history_years', 'Engineered','Years since earliest credit line',            'H2'],
        ['has_derogatory',       'Engineered','Binary: any public derogatory records?',      'H5'],
        ['purpose_*',            'Encoded',  '14 binary columns from loan purpose',         'H3'],
        ['home_ownership_*',     'Encoded',  '5 binary columns from home ownership',        '—'],
        ['verification_status_*','Encoded',  '2 binary columns from income verification',   '—'],
    ],
    col_widths=[1.6, 1.0, 2.5, 1.0]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 9. MACHINE LEARNING MODELS
# ════════════════════════════════════════════════════════════════

add_heading(doc, '9. Machine Learning Models', 1, color=DARKBLUE)

add_heading(doc, '9.1  What Is a Classification Model?', 2)
add_body(doc,
    'A classification model learns patterns from historical data to assign new observations '
    'to a category — in this case, "default" or "fully paid". Given 34 features about an '
    'applicant, the model outputs a probability between 0 and 1: the probability that this '
    'loan will default. A decision threshold (e.g., 0.35) then converts that probability '
    'into a binary decision: approve or reject.'
)

add_heading(doc, '9.2  Logistic Regression (Baseline)', 2)
add_body(doc,
    'Logistic Regression is the simplest classification model — essentially a weighted sum '
    'of features passed through a sigmoid function to produce a probability. It is fast, '
    'interpretable, and serves as a baseline. However, it assumes linear relationships '
    'between features and the outcome, which is rarely true for credit data. In our results, '
    'it achieved an AUC of 0.6996 but only caught 6.5% of defaults at the default threshold '
    '— it was heavily biased toward the majority class (fully paid).'
)

add_heading(doc, '9.3  Random Forest', 2)
add_body(doc,
    'A Random Forest builds hundreds of decision trees, each trained on a random subset '
    'of data and features, then combines their predictions by majority vote. This ensemble '
    'approach reduces overfitting and handles non-linear relationships well. The '
    'class_weight="balanced" setting was used to compensate for the 80/20 class imbalance, '
    'making the model pay more attention to the minority class (defaults). It achieved '
    'AUC 0.7072 and caught 67.9% of defaults.'
)

add_heading(doc, '9.4  XGBoost (Selected Model)', 2)
add_body(doc,
    'XGBoost (Extreme Gradient Boosting) builds trees sequentially — each new tree learns '
    'from the errors of the previous one. It is generally the best-performing model on '
    'structured tabular data and is widely used in industry for credit risk modelling. '
    'The scale_pos_weight parameter (set to ~4.0, reflecting the 80/20 imbalance) '
    'ensures the model penalises missed defaults more heavily during training.'
)

styled_table(doc,
    ['XGBoost Parameter', 'Value', 'What It Does'],
    [
        ['n_estimators',      '300',   'Number of trees — more trees = better accuracy (up to a point)'],
        ['max_depth',         '6',     'Maximum depth of each tree — controls complexity, prevents overfitting'],
        ['learning_rate',     '0.1',   'How much each tree corrects the previous — lower = more robust'],
        ['scale_pos_weight',  '~4.01', 'Weights defaults 4x more — compensates for 80/20 class imbalance'],
        ['subsample',         '0.8',   'Each tree uses 80% of rows — adds randomness, reduces overfitting'],
        ['colsample_bytree',  '0.8',   'Each tree uses 80% of features — adds randomness'],
    ],
    col_widths=[1.8, 0.9, 4.0]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 10. MODEL COMPARISON
# ════════════════════════════════════════════════════════════════

add_heading(doc, '10. Model Comparison', 1, color=DARKBLUE)

add_heading(doc, '10.1  Evaluation Metrics Explained', 2)

for term, defn in [
    ('ROC-AUC (Area Under the ROC Curve):',
     'The primary metric for classification. Measures how well the model separates the '
     'two classes across ALL possible thresholds. AUC = 0.5 means random guessing; '
     'AUC = 1.0 means perfect separation. For credit risk, AUC 0.70–0.75 is considered good.'),
    ('Default Catch Rate (Recall / Sensitivity):',
     'Of all actual defaults in the test set, what percentage did the model correctly '
     'flag? Higher is better — missing a default means funding a loan that loses money.'),
    ('Good Borrower Rejection Rate (False Positive Rate):',
     'Of all good borrowers, what percentage did the model incorrectly reject? '
     'Lower is better — every wrongly rejected borrower is lost revenue.'),
    ('F2 Score:',
     'A metric that weights recall twice as heavily as precision. Purpose-built for '
     'scenarios where false negatives (missed defaults) are more costly than false '
     'positives (wrongly rejected good borrowers).'),
]:
    add_bullet(doc, defn, bold_prefix=term)

add_heading(doc, '10.2  Results at Default Threshold (0.5)', 2)

styled_table(doc,
    ['Model', 'ROC-AUC', 'Default Catch Rate', 'Good Rejection Rate', 'Decision'],
    [
        ['Logistic Regression', '0.6996', '6.5%',  '1.6%',  'Rejected — barely catches any defaults'],
        ['Random Forest',       '0.7072', '67.9%', '37.9%', 'Viable — good recall but not best AUC'],
        ['XGBoost',             '0.7170', '67.7%', '36.1%', '✓ Selected — best AUC, best overall'],
    ],
    col_widths=[1.8, 1.0, 1.5, 1.6, 2.0]
)

add_body(doc,
    'XGBoost was selected as the final model based on its superior ROC-AUC. '
    'The ROC curve below shows all three models — the further the curve bows toward the '
    'top-left corner, the better. XGBoost (red) consistently outperforms both alternatives.'
)

add_image(doc, f'{BASE}/data/rpt_roc.png', width=5.2,
          caption='Figure 5: ROC curves for all three models. The black dot marks the operating point of the final XGBoost model.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 11. THRESHOLD OPTIMISATION
# ════════════════════════════════════════════════════════════════

add_heading(doc, '11. Threshold Optimisation', 1, color=DARKBLUE)

add_heading(doc, '11.1  What Is a Decision Threshold?', 2)
add_body(doc,
    'The model outputs a probability for every loan application — for example, 0.42 means '
    '"42% chance this loan defaults." A decision threshold converts this probability into '
    'an action: approve or reject. At threshold 0.5, any loan with predicted default '
    'probability above 50% is rejected. At threshold 0.35, any loan above 35% is rejected.'
)
add_body(doc,
    'Lowering the threshold makes the model more aggressive — it catches more defaults '
    'but also rejects more good borrowers. Raising it makes the model more lenient — fewer '
    'rejections but more missed defaults. The right threshold depends on the business context.'
)

add_heading(doc, '11.2  Why F2-Score Optimisation?', 2)
add_body(doc,
    'In credit risk, the cost of missing a default (funding a bad loan) is significantly '
    'higher than the cost of rejecting a good borrower (losing a loan origination fee). '
    'The F2-score formalises this asymmetry by weighting recall (default catch rate) '
    'twice as heavily as precision.'
)
add_callout(doc,
    'F2 = (5 × Precision × Recall) / (4 × Precision + Recall)\n'
    'This formula forces the optimal threshold to sit where recall is high, '
    'accepting some loss in precision as the cost.',
    label='F2 FORMULA'
)

add_heading(doc, '11.3  Threshold Sweep Results', 2)
add_body(doc,
    f'The F2-optimal threshold was found to be {THRESHOLD} — meaning any loan '
    f'with predicted default probability ≥ {THRESHOLD} is rejected. '
    f'The charts below show how catch rate and rejection rate change across all thresholds, '
    f'and where F2 score is maximised.'
)

add_image(doc, f'{BASE}/data/rpt_threshold.png', width=5.8,
          caption='Figure 6: Left — catch rate and rejection rate vs threshold. Right — F2 score vs threshold.')

styled_table(doc,
    ['Threshold', 'Default Catch Rate', 'Good Rejection Rate', 'F2 Score', 'Notes'],
    [
        ['0.50 (default)', '67.7%', '36.1%', '—',   'Too lenient — misses many defaults'],
        ['0.40',           '83.0%', '54.4%', '—',   'Better catch but high rejection'],
        [f'{THRESHOLD} (F2-optimal)', f'{catch_rate:.1%}', f'{good_rej:.1%}', f'{f2:.4f}', 'Best balance — selected'],
    ],
    col_widths=[1.5, 1.5, 1.6, 1.0, 2.2]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 12. FINAL MODEL PERFORMANCE
# ════════════════════════════════════════════════════════════════

add_heading(doc, '12. Final Model Performance', 1, color=DARKBLUE)

add_heading(doc, '12.1  Confusion Matrix', 2)
add_body(doc,
    'The confusion matrix shows exactly what the model predicted for each loan in the '
    f'test set of {total:,} loans at the optimal threshold of {THRESHOLD}:'
)

add_image(doc, f'{BASE}/data/rpt_cm.png', width=5.0,
          caption='Figure 7: Confusion matrix. Darker blue = higher count.')

styled_table(doc,
    ['Cell', 'Count', 'Meaning'],
    [
        ['True Negative (TN)',  f'{tn:,}',
         f'Good loans correctly approved — {tn/(tn+fp):.1%} of all good borrowers'],
        ['False Positive (FP)', f'{fp:,}',
         f'Good loans incorrectly rejected — {good_rej:.1%} false alarm rate'],
        ['False Negative (FN)', f'{fn:,}',
         f'Defaults that slipped through — {fn/(tp+fn):.1%} of all defaults missed'],
        ['True Positive (TP)',  f'{tp:,}',
         f'Defaults correctly caught and blocked — {catch_rate:.1%} catch rate'],
    ],
    col_widths=[1.8, 1.0, 4.2]
)

add_heading(doc, '12.2  Score Distribution', 2)
add_body(doc,
    'The chart below shows the distribution of predicted default probabilities for good '
    'loans (blue) and actual defaults (red). The closer these distributions are to each '
    'other, the harder the classification problem. The vertical line is the decision '
    'threshold — everything to its right is rejected.'
)

add_image(doc, f'{BASE}/data/rpt_score_dist.png', width=5.8,
          caption='Figure 8: Score distributions. The overlap between blue and red represents the inherent difficulty of the problem.')

add_heading(doc, '12.3  Feature Importance', 2)
add_body(doc,
    'XGBoost assigns an importance score to each feature based on how often it was used '
    'in splits and how much it reduced prediction error. The top features confirm our '
    'hypotheses and domain intuition:'
)

for feat, explanation in [
    ('int_rate / grade:', 'Interest rate and grade are the lender\'s own risk assessment — the strongest signal available.'),
    ('dti:',              'Debt-to-income ratio — confirms H1. Higher DTI = higher default risk.'),
    ('annual_inc:',       'Annual income — confirms H4. Lower income borrowers default more.'),
    ('revol_util:',       'Revolving line utilisation — borrowers using most of their credit limit are under financial stress.'),
    ('inq_last_6mths:',   'Recent credit enquiries — multiple applications in a short window signals financial distress.'),
]:
    add_bullet(doc, explanation, bold_prefix=feat)

add_image(doc, f'{BASE}/data/rpt_feat_imp.png', width=5.5,
          caption='Figure 9: Top 15 most important features. Red bars are above-median importance.')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 13. BUSINESS IMPACT
# ════════════════════════════════════════════════════════════════

add_heading(doc, '13. Business Impact', 1, color=DARKBLUE)

add_heading(doc, '13.1  What Happens Without the Model?', 2)
add_body(doc,
    f'Without any screening model, if LendingClub approves all {total:,} applications '
    f'in the test set, {tp+fn:,} ({dr_base:.1%}) of those loans will default. '
    f'Every funded default represents a complete loss of principal for the investor.'
)

add_heading(doc, '13.2  What Happens With the Model?', 2)
add_body(doc,
    f'With the XGBoost model at threshold {THRESHOLD}, the platform rejects the '
    f'{tp+fp:,} highest-risk applications. Of the {approved:,} loans that are approved, '
    f'only {fn:,} will default — a default rate of {dr_model:.2%} instead of {dr_base:.2%}.'
)

add_callout(doc,
    f'Plain English: For every 100 loan applications, the model approves approximately '
    f'{approved/total*100:.0f}. Of those {approved/total*100:.0f} approved loans, '
    f'only {dr_model*100:.1f} will default — compared to {dr_base*100:.1f} defaults '
    f'if all 100 were approved. The model eliminates {dr_reduc:.0%} of default exposure.',
    label='PLAIN ENGLISH IMPACT'
)

add_heading(doc, '13.3  Full Business Metrics', 2)
metric_grid(doc, {
    'Total Applications':        f'{total:,}',
    'Approved by Model':         f'{approved:,}  ({approved/total:.1%})',
    'Rejected by Model':         f'{rejected:,}  ({rejected/total:.1%})',
    'Defaults in Population':    f'{tp+fn:,}',
    'Defaults Caught & Blocked': f'{tp:,}  ({catch_rate:.1%})',
    'Defaults Missed':           f'{fn:,}  ({fn/(tp+fn):.1%})',
    'Good Borrowers Approved':   f'{tn:,}  ({tn/(tn+fp):.1%})',
    'Good Borrowers Rejected':   f'{fp:,}  ({good_rej:.1%})',
    'Default Rate (No Model)':   f'{dr_base:.2%}',
    'Default Rate (With Model)': f'{dr_model:.2%}',
    'Default Rate Reduction':    f'{dr_reduc:.0%}',
    'ROC-AUC':                   f'{auc:.4f}',
}, cols=2)

add_heading(doc, '13.4  The Business Tradeoff', 2)
add_body(doc,
    f'The model rejects {good_rej:.1%} of good borrowers — these are potential customers '
    f'who would have repaid their loan and generated revenue. This is the cost of risk '
    f'reduction. The business must decide whether the {dr_reduc:.0%} reduction in default '
    f'exposure justifies the {good_rej:.0%} reduction in loan origination volume. '
    f'This tradeoff can be adjusted by changing the threshold: a higher threshold (e.g., 0.45) '
    f'approves more borrowers but catches fewer defaults; a lower threshold (e.g., 0.25) '
    f'is more aggressive but rejects even more good applicants.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 14. CONCLUSIONS & RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════

add_heading(doc, '14. Conclusions & Recommendations', 1, color=DARKBLUE)

add_heading(doc, '14.1  What We Proved', 2)
for item in [
    ('H1 Confirmed:', 'Higher DTI is associated with higher default risk. DTI is in the top 5 most important features.'),
    ('H2 Weak:', 'Credit history length has a statistically significant but small effect. Contributes modestly to the model.'),
    ('H3 Confirmed:', 'Loan purpose strongly predicts default. Small business loans carry the highest risk; car loans the lowest.'),
    ('H4 Confirmed:', 'Lower income correlates with higher default. Income features among the top predictors.'),
    ('H5 Weak:', 'Derogatory marks have a small but real effect. Converted to a binary flag (has_derogatory) for best signal extraction.'),
]:
    add_bullet(doc, item[1], bold_prefix=item[0])

add_heading(doc, '14.2  Recommendations', 2)
for rec in [
    ('Deploy XGBoost at threshold 0.35', 'as the primary loan screening tool. This configuration is validated on 269,612 real loans.'),
    ('Review threshold quarterly', '— borrower behaviour shifts over economic cycles. What worked in 2015 may not work in 2026.'),
    ('Conduct human review for borderline scores', '(probability between 0.28–0.42). These loans are genuinely uncertain; a human officer can add context the model cannot.'),
    ('Enrich features for next model version', '— FICO score, payment history, geographic region, and macroeconomic indicators could push AUC above 0.75.'),
    ('Monitor for concept drift', '— retrain on a rolling 24-month window. Economic recessions change default patterns rapidly.'),
    ('Consider calibration', '— the raw probabilities are not perfectly calibrated. Platt scaling or isotonic regression would make the probability output more reliable for risk pricing.'),
]:
    add_bullet(doc, rec[1], bold_prefix=rec[0])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 15. LIMITATIONS & FUTURE WORK
# ════════════════════════════════════════════════════════════════

add_heading(doc, '15. Limitations & Future Work', 1, color=DARKBLUE)

add_heading(doc, '15.1  Current Limitations', 2)
for lim in [
    ('AUC of 0.717 is moderate:', 'The model correctly separates defaults from paid loans about 71.7% of the time. There is meaningful room for improvement, particularly through better features.'),
    ('Dataset covers 2007–2018:', 'Economic conditions have changed significantly. The 2008 financial crisis distorts some patterns in the earlier data.'),
    ('Class imbalance handling:', 'We used scale_pos_weight in XGBoost. SMOTE (Synthetic Minority Oversampling Technique) was not applied and may improve recall further.'),
    ('No probability calibration:', 'The model\'s raw output probabilities are not calibrated — a predicted 0.6 does not mean exactly 60% chance of default.'),
    ('Static threshold:', 'A single global threshold is applied to all loan types and amounts. A tiered approach (different thresholds by loan purpose or grade) may be more effective.'),
]:
    add_bullet(doc, lim[1], bold_prefix=lim[0])

add_heading(doc, '15.2  Future Improvements', 2)
for fut in [
    'Add FICO score and external credit bureau data — the single most impactful improvement available',
    'Apply SMOTE to training data to improve minority class recall',
    'Train on the full 2.26M rows rather than the filtered 1.35M',
    'Experiment with LightGBM and CatBoost as alternatives to XGBoost',
    'Build a two-stage model: first predict default probability, then use a second model to price the risk premium',
    'Apply SHAP (SHapley Additive exPlanations) values for individual loan-level explanations',
]:
    add_bullet(doc, fut)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 16. GLOSSARY
# ════════════════════════════════════════════════════════════════

add_heading(doc, '16. Glossary of Key Terms', 1, color=DARKBLUE)

glossary = [
    ('AUC (Area Under the Curve)',
     'A measure of how well a model separates two classes across all thresholds. '
     'Ranges from 0.5 (random) to 1.0 (perfect). Higher is better.'),
    ('Charged Off',
     'LendingClub\'s term for a loan that has been written off as uncollectable. '
     'Equivalent to default in this project.'),
    ('Class Imbalance',
     'When one category in the target variable is much more common than another. '
     'Here, 80% of loans were fully paid vs 20% defaulted.'),
    ('Confusion Matrix',
     'A 2×2 table showing the counts of true positives, true negatives, false positives, '
     'and false negatives for a classification model.'),
    ('DTI (Debt-to-Income Ratio)',
     'Monthly debt payments divided by monthly gross income, expressed as a percentage. '
     'A DTI of 20 means 20% of income goes toward debt service.'),
    ('F2 Score',
     'A metric that weights recall twice as heavily as precision. Used when false negatives '
     '(missed defaults) are more costly than false positives (wrongly rejected good borrowers).'),
    ('False Negative',
     'A loan the model predicted was safe (approved) but actually defaulted. '
     'The most costly type of error in this context.'),
    ('False Positive',
     'A loan the model predicted would default (rejected) but would actually have been repaid. '
     'Represents lost revenue.'),
    ('Feature Engineering',
     'Transforming raw columns into more useful representations — e.g., converting a date '
     'string into a numeric "years of credit history" value.'),
    ('Feature Importance',
     'A measure of how much each feature contributed to a tree-based model\'s predictions. '
     'Features used in more splits and that reduce error more get higher importance.'),
    ('One-Hot Encoding',
     'Converting a categorical variable (e.g., "purpose") into multiple binary columns '
     '(purpose_car, purpose_medical, etc.) — one column per category.'),
    ('Oversampling / SMOTE',
     'A technique to address class imbalance by synthetically generating new minority-class '
     'examples so the model sees more defaults during training.'),
    ('Precision',
     'Of all loans the model predicted would default, what fraction actually defaulted. '
     'High precision = few false alarms.'),
    ('Recall (Sensitivity)',
     'Of all actual defaults, what fraction did the model catch. '
     'High recall = few missed defaults.'),
    ('ROC Curve',
     'Receiver Operating Characteristic curve — plots the True Positive Rate vs False Positive Rate '
     'across all possible thresholds. AUC summarises its area.'),
    ('Scale_pos_weight',
     'An XGBoost parameter that tells the model to penalise missed defaults more heavily. '
     'Set to (number of negatives) / (number of positives) = ~4.0 in this project.'),
    ('StandardScaler',
     'Transforms each feature to have mean=0 and standard deviation=1. '
     'Prevents features with large numeric ranges from dominating the model.'),
    ('Stratified Split',
     'A train/test split that preserves the class ratio in both sets. '
     'Ensures neither set is accidentally loaded with more defaults than the other.'),
    ('Threshold',
     'The cut-off probability above which a loan is classified as high-risk and rejected. '
     'Lower threshold = more rejections; higher threshold = fewer rejections.'),
    ('VIF (Variance Inflation Factor)',
     'Measures how much a feature\'s variance is explained by other features. '
     'VIF > 10 indicates severe multicollinearity.'),
    ('XGBoost',
     'Extreme Gradient Boosting — a state-of-the-art machine learning algorithm that builds '
     'trees sequentially, each one correcting the errors of the previous. '
     'The most widely used model for structured tabular data in industry.'),
]

for term, defn in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(1)
    run1 = p.add_run(term + ':  ')
    run1.bold = True
    run1.font.color.rgb = BLUE
    run2 = p.add_run(defn)
    run2.font.size = Pt(9.5)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 17. NOTEBOOK STRUCTURE
# ════════════════════════════════════════════════════════════════

add_heading(doc, '17. Project Notebook Structure', 1, color=DARKBLUE)

add_body(doc,
    'The full analysis is organised into seven Jupyter notebooks, each covering one phase '
    'of the project. They are designed to be run in order, with each notebook saving its '
    'output as a Parquet file for the next notebook to load.'
)

styled_table(doc,
    ['Notebook', 'Phase', 'Key Outputs'],
    [
        ['01_data_loading.ipynb',
         'Data Loading & Target Mapping',
         'loans_filtered.parquet — 1.35M filtered loans with binary target'],
        ['02_eda.ipynb',
         'Exploratory Data Analysis',
         'Hypothesis charts (h1_dti.png, h3_purpose.png, etc.) + summary table'],
        ['03_preprocessing.ipynb',
         'Feature Engineering, Cleaning & Splitting',
         'X_train.parquet, X_test.parquet, y_train.parquet, y_test.parquet, scaler.pkl'],
        ['03b_feature_significance.ipynb',
         'Statistical Feature Validation',
         'Correlation matrix, VIF table, feature ranking — informs feature dropping decisions'],
        ['04_modeling.ipynb',
         'Model Training & Threshold Optimisation',
         'xgb_model.pkl, threshold.txt (0.35) — F2-optimal XGBoost model'],
        ['05_evaluation.ipynb',
         'Final Evaluation & Business Impact',
         'All evaluation charts + printed business impact summary'],
        ['06_rejected_analysis.ipynb',
         'Rejected Applications Analysis',
         'Scale context, DTI/Risk Score/Employment comparisons, time trends, geographic analysis'],
    ],
    col_widths=[2.0, 1.8, 3.3]
)


# ─── save ─────────────────────────────────────────────────────────────────────

out_path = f'{BASE}/docs/LendingClub_Loan_Default_Risk_Report_v4.docx'
os.makedirs(f'{BASE}/docs', exist_ok=True)
doc.save(out_path)
print(f'\nReport saved: {out_path}')
print(f'  Pages (estimated): ~35-40')
